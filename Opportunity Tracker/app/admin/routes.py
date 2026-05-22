from flask import Blueprint,url_for,render_template,redirect,flash,get_flashed_messages,request,session,abort
from shared_models.models import Authentication,Employee,Administrator,EmployeeEmail,OpportunityTracker,Comment
from shared_models import db
from sqlalchemy.exc import IntegrityError
from app.utils import login_required,admin_required
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


admin_bp = Blueprint("OTadmin",__name__,url_prefix="opportunity_tracker/admin")

opportunity_stages = {"Lead":"orange","RFP":"yellow","Presentation":"yellow","Demo":"yellow","Contract":"yellowgreen","Won":"green","Lost":"red"}
modules = ["ERP","HCM","SCM","EPM","PaaS","Primaver","Payroll"]

@admin_bp.route("/",methods=["GET"])
@login_required
@admin_required
def home():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("home.html",state="admin",role="admin",auth_name=auth_name,auth=auth)

@admin_bp.route("/logout",methods=["POST","GET"])
@login_required
@admin_required
def logout():
    if request.method == "POST":
        session.clear()
        return redirect(url_for("base.login"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("logout.html",state="admin",role="admin",auth=auth)

@admin_bp.route("/select_opportunity_tracker",methods=["POST","GET"])
@login_required
@admin_required
def select_opportunity_tracker():
    batches = OpportunityTracker.query.all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("select_opportunity_tracker.html",auth=auth,role="admin",batches=batches,state="admin")

@admin_bp.route("/reports",methods=["POST","GET"])
@login_required
@admin_required
def reports():
    auth = Authentication.query.get(session.get("user_id"))
    opportunity_trackers = OpportunityTracker.query.all()
    now = datetime.now()
    return render_template("reports.html",auth=auth,role="admin",state="admin",reports=reports,now=now,opportunity_trackers=opportunity_trackers,opportunity_stages=opportunity_stages)

@admin_bp.route("/add_member", methods=["POST", "GET"])
@login_required
@admin_required
def add_member():
    auth_name = Authentication.query.get(session["user_id"]).name
    if request.method == "POST":
        email = request.form.get("email")
        department = request.form.get("department")

        existing_email = EmployeeEmail.query.filter_by(email=email).first()

        if existing_email and existing_email.department != department:
            flash("This email is already enrolled in another department.", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

        if existing_email and existing_email.department == department:
            flash(f"This email is already enrolled in this department.", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

        if not existing_email:
            new_entry = EmployeeEmail(email=email, department=department)
            db.session.add(new_entry)
            print(1)
            db.session.commit()
            flash("Member enrolled successfully.", "success")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(url_for("admin.home"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

@admin_bp.route("/add_opportunity_tracker",methods=["POST","GET"])
@login_required
@admin_required
def add_opportunity_tracker():
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    print(names_auth)
    if request.method == "POST":
        year = request.form.get("year")
        firm = request.form.get("firm")
        status = request.form.get("status")
        industry = request.form.get("industry")
        modl = request.form.getlist("modules[]")
        reffered_by = request.form.get("refferred_by")
        deal_size = request.form.get("deal_size")
        opportunity_stage = request.form.get("opportunity_stage")
        exp_closure_date = request.form.get("exp_closure_date")
        customer_party = request.form.get("customer_party")
        assigned_to_ids = request.form.getlist("assigned_tos[]")
        print(assigned_to_ids)
        assigned_tos = Authentication.query.filter(Authentication.id.in_(assigned_to_ids)).all()
        now = datetime.now()
        if OpportunityTracker.query.filter_by(firm=firm).first():
            flash(f"Opportunity Tracker for {firm} already exists ","error")
            return redirect(url_for("admin.opportunity_trackers"))
        new_rm = OpportunityTracker(firm=firm, year=year, status=status, industry=industry, modules=modl, reffered_by=reffered_by, deal_size=deal_size, opportunity_stage=opportunity_stage, exp_closure_date=exp_closure_date, customer_party=customer_party,assigned_to=assigned_tos)
        if new_rm:
            db.session.add(new_rm)
        db.session.commit()
        flash("Opportunity Tracker created successfully","success")
        return redirect(url_for("admin.opportunity_trackers"))
    return render_template("add_opportunity_tracker.html",auth=auth,role="admin",state="admin",names_auth=names_auth,opportunity_stages=opportunity_stages,modules=modules)

@admin_bp.route("/edit_opportunity_tracker/<int:rm_id>", methods=["POST", "GET"])
@login_required
@admin_required
def edit_opportunity_tracker(rm_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    participants = OpportunityTracker.query.get(rm_id).assigned_to
    participant_ids = [user.id for user in participants]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    opportunity_tracker = OpportunityTracker.query.get(rm_id)
    now = datetime.now()
    print(opportunity_tracker.assigned_to)
    print(names_auth[0])
    if request.method == "POST":
        year = request.form.get("year")
        status = request.form.get("status")
        industry = request.form.get("industry")
        modls = request.form.getlist("modules[]")
        refferred_by = request.form.get("refferred_by")
        deal_size = request.form.get("deal_size")
        opportunity_stage = request.form.get("opportunity_stage")
        exp_closure_date = request.form.get("exp_closure_date")
        customer_party = request.form.get("customer_party")
        assigned_to_ids = request.form.getlist("assigned_tos[]")
        print(assigned_to_ids)
        print(deal_size)
        assigned_tos = Authentication.query.filter(Authentication.id.in_(assigned_to_ids)).all()
        opportunity_tracker.status = status
        opportunity_tracker.industry = industry
        opportunity_tracker.modules = modls
        opportunity_tracker.reffered_by = refferred_by
        opportunity_tracker.deal_size = deal_size
        opportunity_tracker.opportunity_stage = opportunity_stage
        opportunity_tracker.exp_closure_date = exp_closure_date
        opportunity_tracker.customer_party = customer_party
        opportunity_tracker.assigned_to = assigned_tos
        db.session.commit()
        flash("Opportunity Tracker Edited successfully","success")
        now = datetime.now()
        return redirect(url_for("admin.opportunity_trackers"))
    return render_template("edit_opportunity_tracker.html",Title="EDIT OBJECTIVES",state='admin',auth=auth,role="admin",now=now,names_auth=names_auth,opportunity_stages=opportunity_stages,modules=modules,opportunity_tracker=opportunity_tracker)

@admin_bp.route("/opportunity_trackers", methods=["POST","GET"])
@login_required
@admin_required
def opportunity_trackers():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth =  Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    opportunity_trackers = OpportunityTracker.query.all()
    return render_template("opportunity_trackers.html",auth=auth,role="admin",mode="See All",now=now,opportunity_trackers=opportunity_trackers,state="admin",opportunity_stages=opportunity_stages,modules=modules)

@admin_bp.route("/opportunity_tracker_overview/<int:rm_id>", methods=["POST", "GET"])
@login_required
@admin_required
def opportunity_tracker_overview(rm_id):
    opportunity_tracker = OpportunityTracker.query.get(rm_id)
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    if request.method == "POST":
        comment_text = request.form.get("comment_text", "").strip()
        if comment_text:
            new_comment = Comment(
                opportunity_tracker_id=rm_id,
                author=auth.name,
                comment=comment_text,
                timestamp=datetime.now(),
            )
            db.session.add(new_comment)
            db.session.commit()

        return redirect(url_for("admin.opportunity_tracker_overview", rm_id=rm_id))  # reload page
    comments = Comment.query.filter_by(opportunity_tracker_id=rm_id).order_by(Comment.timestamp).all()
    return render_template("opportunity_tracker_overview.html",state='admin',auth=auth,role="admin",opportunity_tracker=opportunity_tracker,now=now,comments=comments,opportunity_stages=opportunity_stages,modules=modules)

@admin_bp.route("/download-report-word/<int:objective_id>")
def download_report_word(objective_id):
    doc = Document()
    doc.add_heading("Objective Report", 0)
    doc.add_paragraph(f"Objective ID: {objective_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Message"
    hdr[1].text = "Status"
    hdr[2].text = "Timestamp"

    # add data rows

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@admin_bp.route("/download-report/<int:objective_id>")
def download_report(objective_id):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # write content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Objective Report")

    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, f"Objective ID: {objective_id}")
    p.drawString(50, height - 100, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p.save()
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")
