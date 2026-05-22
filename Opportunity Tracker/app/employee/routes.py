from flask import Blueprint,render_template,request,url_for,redirect,session,flash,get_flashed_messages,abort
from app.utils import login_required,employee_required
from shared_models.models import Authentication,Employee,Comment,OpportunityTracker
from shared_models import db
from sqlalchemy.exc import IntegrityError
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


employee_bp = Blueprint("OTemployee",__name__,url_prefix="opportunity_tracker/employee")

opportunity_stages = {"Lead":"orange","RFP":"yellow","Presentation":"yellow","Demo":"yellow","Contract":"yellowgreen","Won":"green","Lost":"red"}
modules = ["ERP","HCM","SCM","EPM","PaaS","Primaver","Payroll"]


@employee_bp.route("/")
@login_required
@employee_required
def home():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("home.html",state="employee",role="employee",auth_name=auth_name,auth=auth)

@employee_bp.route("/logout",methods=["POST","GET"])
@login_required
@employee_required
def logout():
    auth_name = Authentication.query.get(session["user_id"]).name
    if request.method == "POST":
        session.pop("role",None)
        session.pop("user_id",None)
        session.pop("email",None)
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("base.login"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("logout.html",state="employee",role="employee",auth=auth)

@employee_bp.route("/select_opportunity_tracker",methods=["POST","GET"])
@login_required
@employee_required
def select_opportunity_tracker():
    auth = Authentication.query.get(session.get("user_id"))
    opportunity_trackers = OpportunityTracker.query.filter(OpportunityTracker.assigned_to.any(Authentication.id==auth.id))
    return render_template("select_opportunity_tracker.html",auth=auth,role="employee",opportunity_trackers=opportunity_trackers,state="employee")


@employee_bp.route("/reports",methods=["POST","GET"])
@login_required
@employee_required
def reports():
    auth = Authentication.query.get(session.get("user_id"))
    opportunity_trackers = OpportunityTracker.query.filter(OpportunityTracker.assigned_to.any(Authentication.id == auth.id)).all()
    now = datetime.now()
    return render_template("reports.html",auth=auth,role="employee",state="employee",reports=reports,now=now,opportunity_trackers=opportunity_trackers)

@employee_bp.route("/opportunity_trackers", methods=["POST","GET"])
@login_required
@employee_required
def opportunity_trackers():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth =  Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    opportunity_trackers = OpportunityTracker.query.filter(OpportunityTracker.assigned_to.any(Authentication.id == auth.id)).all()
    return render_template("opportunity_trackers.html",auth=auth,role="employee",mode="See All",now=now,opportunity_trackers=opportunity_trackers,state="employee",modules=modules,opportunity_stages=opportunity_stages)

@employee_bp.route("/opportunity_tracker_overview/<int:rm_id>", methods=["POST", "GET"])
@login_required
@employee_required
def opportunity_tracker_overview(rm_id):
    opportunity_tracker = OpportunityTracker.query.get(rm_id)
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    if request.method == "POST":
        comment_text = request.form.get("comment_text", "").strip()
        if comment_text:
            new_comment = Comment(
                opportunity_tracker_id=rm_id,
                author=auth.name,
                comment=comment_text,
                timestamp=datetime.now(),
            )
            db.session.add(new_comment)
            db.session.commit()

        return redirect(url_for("employee.opportunity_tracker_overview", rm_id=rm_id))  # reload page

    # GET
    comments = Comment.query.filter_by(opportunity_tracker_id=rm_id).order_by(Comment.timestamp).all()
    return render_template("opportunity_tracker_overview.html",state='employee',auth=auth,role="employee",opportunity_tracker=opportunity_tracker,now=now,comments=comments,opportunity_stages=opportunity_stages,modules=modules)

@employee_bp.route("/edit_opportunity_tracker/<int:rm_id>", methods=["POST", "GET"])
@login_required
@employee_required
def edit_opportunity_tracker(rm_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    participants = OpportunityTracker.query.get(rm_id).assigned_to
    participant_ids = [user.id for user in participants]
    names_auth = Authentication.query.order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    opportunity_tracker = OpportunityTracker.query.get(rm_id)
    now = datetime.now()
    print(opportunity_tracker.assigned_to)
    print(names_auth[0])
    if request.method == "POST":
        industry = request.form.get("industry")
        deal_size = request.form.get("deal_size")
        opportunity_stage = request.form.get("opportunity_stage")   
        customer_party = request.form.get("customer_party")
        opportunity_tracker.industry = industry
        opportunity_tracker.deal_size = deal_size
        opportunity_tracker.opportunity_stage = opportunity_stage
        opportunity_tracker.customer_party = customer_party
        db.session.commit()
        flash("Opportunity Tracker Edited successfully","success")
        now = datetime.now()
        return redirect(url_for("employee.opportunity_trackers"))
    return render_template("edit_opportunity_tracker.html",Title="EDIT OBJECTIVES",state='employee',auth=auth,role="employee",now=now,names_auth=names_auth,opportunity_stages=opportunity_stages,modules=modules,opportunity_tracker=opportunity_tracker)


@employee_bp.route("/download-report-word/<int:objective_id>")
def download_report_word(objective_id):
    doc = Document()
    doc.add_heading("Objective Report", 0)
    doc.add_paragraph(f"Objective ID: {objective_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Message"
    hdr[1].text = "Status"
    hdr[2].text = "Timestamp"

    # add data rows
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@employee_bp.route("/download-report/<int:objective_id>")
def download_report(objective_id):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # write content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Objective Report")

    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, f"Objective ID: {objective_id}")
    p.drawString(50, height - 100, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p.save()
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")
